import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

CHUNK_SIZE = 700
CHUNK_OVERLAP = 100


def decode_document(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    return content.decode("utf-8", errors="ignore")


def extract_document_text(filename: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()

    if extension in {".txt", ".md"}:
        return decode_document(content)
    if extension == ".pdf":
        return extract_pdf_text(content)
    if extension == ".docx":
        return extract_docx_text(content)

    raise ValueError("Unsupported FAQ file type.")


def extract_pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ValueError("Unable to parse PDF text.") from exc

    return "\n\n".join(page.strip() for page in pages if page.strip())


def extract_docx_text(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise ValueError("Unable to parse DOCX text.") from exc

    parts: list[str] = []
    parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    normalized = normalize_text(text)
    if not normalized:
        return []

    paragraphs = [paragraph.strip() for paragraph in normalized.split("\n\n") if paragraph.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        if len(paragraph) > chunk_size:
            if current:
                chunks.append(current)
                current = ""
            chunks.extend(split_long_text(paragraph, chunk_size, overlap))
            continue

        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            chunks.append(current)
            current = paragraph

    if current:
        chunks.append(current)

    return chunks


def split_long_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    chunks: list[str] = []
    start = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == len(text):
            break
        start = max(end - overlap, start + 1)

    return chunks
